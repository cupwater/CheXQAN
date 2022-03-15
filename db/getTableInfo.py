# encoding: utf8
'''
Author: Baoyun Peng
Date: 2022-03-14 10:59:35
LastEditTime: 2022-03-14 13:41:43
Description: 

'''
import argparse
import pymysql
from pymysql import OperationalError
from docx import Document

ReportSql = """
SELECT
	COLUMN_NAME 列名,
	COLUMN_TYPE 数据类型,
	IS_NULLABLE 是否为空,
	COLUMN_KEY 约束条件,
	COLUMN_COMMENT 说明
FROM
    INFORMATION_SCHEMA.COLUMNS 
WHERE
    table_schema = '{}' 
AND table_name = '{}'
"""


class DB(object):

    def __init__(self, database):
        self.setting = {
            "host": '127.0.0.1',
            "port": 3306,
            "user": 'download',
            "password": 'Down@0221',
            "database": database,
        }
        # 创建数据库连接
        self.dbconn = pymysql.connect(**self.setting, local_infile=1)
        while True:
            try:
                self.dbconn.ping()
                break
            except OperationalError:
                self.dbconn.ping(True)
        # 创建字典型游标(返回的数据是字典类型)
        self.dbcur = self.dbconn.cursor(cursor=pymysql.cursors.DictCursor)

    # __enter__() 和 __exit__() 是with关键字调用的必须方法
    # with本质上就是调用对象的enter和exit方法
    def __enter__(self):
        # 返回游标
        return self.dbcur

    def __exit__(self, exc_type, exc_value, exc_trace):
        # 提交事务
        self.dbconn.commit()

        # 关闭游标
        self.dbcur.close()

        # 关闭数据库连接
        self.dbconn.close()


def Tables(cou, database):
    sql = "select DISTINCT table_name from information_schema.columns where table_schema='{}';".format(
        database)
    cou.execute(sql)
    res = cou.fetchall()
    res = list(map(lambda x: x["table_name"], res))
    return res


def lay_data(database, result):
    document = Document()
    for key in result:
        for tab,data in key.items():
            table = document.add_table(rows=len(data) + 2, cols=5)
            table.style = 'Table Grid'
            table.rows[0].cells[0].text = "表名"
            coll_1 = table.cell(0, 1)
            coll_2 = table.cell(0, 4)
            coll_1.merge(coll_2)
            table.rows[0].cells[1].text = tab
            table.rows[1].cells[0].text = '列名'
            table.rows[1].cells[1].text = '数据类型'
            table.rows[1].cells[2].text = '是否为空'
            table.rows[1].cells[3].text = '约束条件'
            table.rows[1].cells[4].text = '说明'

            # 制作表格
            for index in range(2, len(data) + 2):
                data_index = index - 2
                table.rows[index].cells[0].text = data[data_index][0]
                table.rows[index].cells[1].text = data[data_index][1]
                table.rows[index].cells[2].text = data[data_index][2]
                table.rows[index].cells[3].text = data[data_index][3]
                table.rows[index].cells[4].text = data[data_index][4]
            document.add_paragraph('\n')
    document.save('report/{}.docx'.format(database))


def main(database):
    with DB(database=database) as _cou:
        tables = Tables(_cou, database)
        data = []
        for table in tables:
            _cou.execute(ReportSql.format(database, table))
            data.append({table: list(map(lambda x: list(x.values()), _cou.fetchall()))})
        # lay_data(database, data)


if __name__ == '__main__':
    main(database='ai_model_quality_control')


